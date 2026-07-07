from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree
import sys
import os

def html_to_docx(html_file, docx_file):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(11)
    font.color.rgb = RGBColor(48, 49, 51)
    style.paragraph_format.line_spacing = 1.8
    style.paragraph_format.space_after = Pt(6)

    with open(html_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    parser = etree.HTMLParser()
    tree = etree.fromstring(content, parser)
    
    content_div = tree.find('.//div[@class="content"]')
    if content_div is None:
        content_div = tree.find('.//body')
    
    toc_processed = False
    
    for element in content_div.iter():
        if element.tag == 'div':
            if 'toc' in element.get('class', ''):
                toc_processed = True
                toc_title = element.find('.//h2')
                if toc_title is not None:
                    p = doc.add_paragraph()
                    run = p.add_run(toc_title.text.strip())
                    run.bold = True
                    run.font.size = Pt(14)
                    run.font.color.rgb = RGBColor(64, 158, 255)
                continue
        
        if element.tag == 'h2':
            text = ''.join(element.itertext()).strip()
            if text:
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(48, 49, 51)
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(8)
        
        elif element.tag == 'h3':
            text = ''.join(element.itertext()).strip()
            if text:
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(64, 158, 255)
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(6)
        
        elif element.tag == 'p':
            text = ''.join(element.itertext()).strip()
            if text:
                p = doc.add_paragraph(text)
                p.paragraph_format.first_line_indent = Cm(0.74)
        
        elif element.tag == 'ul':
            for li in element.findall('.//li'):
                text = ''.join(li.itertext()).strip()
                if text:
                    p = doc.add_paragraph()
                    run = p.add_run('● ' + text)
                    p.paragraph_format.first_line_indent = Cm(0.74)
        
        elif element.tag == 'ol':
            index = 1
            for li in element.findall('.//li'):
                text = ''.join(li.itertext()).strip()
                if text:
                    p = doc.add_paragraph()
                    run = p.add_run(f'{index}. ' + text)
                    p.paragraph_format.first_line_indent = Cm(0.74)
                    index += 1
        
        elif element.tag == 'table':
            rows = element.findall('.//tr')
            if rows:
                first_row = rows[0]
                cols = len(first_row.findall('.//th'))
                if cols == 0:
                    cols = len(first_row.findall('.//td'))
                
                table = doc.add_table(rows=len(rows), cols=cols)
                table.style = 'Table Grid'
                
                for i, row in enumerate(rows):
                    cells = row.findall('.//th')
                    if not cells:
                        cells = row.findall('.//td')
                    
                    for j, cell in enumerate(cells):
                        text = ''.join(cell.itertext()).strip()
                        table.cell(i, j).text = text
                        if i == 0:
                            for paragraph in table.cell(i, j).paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
        
        elif element.tag == 'hr':
            doc.add_page_break()
    
    doc.save(docx_file)
    print(f"Successfully converted {html_file} to {docx_file}")

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python html_to_docx.py <input_html> <output_docx>")
        sys.exit(1)
    
    html_file = sys.argv[1]
    docx_file = sys.argv[2]
    
    if not os.path.exists(html_file):
        print(f"File not found: {html_file}")
        sys.exit(1)
    
    html_to_docx(html_file, docx_file)